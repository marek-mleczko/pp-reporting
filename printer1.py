# -*- coding: utf-8 -*-
"""
Created on Fri Mar 22 14:14:59 2019

@author: marek-mleczko
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Jan 18 15:44:09 2019

@author: marek-mleczko
"""
import time
import xlwt
import docx
import liczenie3 as ex
from docx import Document
import win32com.client as win32
msword = win32.gencache.EnsureDispatch('Word.Application')

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
#shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
start_time = time.time()
doc = Document()

doc1 = Document("pp.docx")

book = xlwt.Workbook()
sheet1 = book.add_sheet('Sheet 1')
row = 0

styles = doc1.styles
#doc1.add_heading('A New Title for my Document', 0)


paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

#from docx.enum.style import WD_STYLE_TYPE
#from docx.enum.text import WD_TAB_ALIGNMENT
#styles = doc1.styles
#style = styles.add_style("Header", WD_STYLE_TYPE.PARAGRAPH)
#style.base_style = styles["Normal"]
#tab_stops = style.paragraph_format.tab_stops
##tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
#tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)










for style in paragraph_styles:
   print(style.name)

print("tabelki \n \n ")
   
for style in table_styles:
   print(style.name)






paragraph1 = doc1.add_paragraph('Lorem ipsum dolor sit amet.')
doc1.add_picture('globus.png')
doc1.add_heading('1.	ORGANIZATION REMUNERATION POLICY')
paragraph = doc1.add_paragraph(style ="Normal")

def drukuj_main(pytanie1):
    if len(pytanie1.sekcja) >1:
            
            doc1.add_paragraph(style ="Normal")
            doc1.add_paragraph(pytanie1.sekcja[1], style="Heading 1")
    else:     
       pass 
    parag = doc1.add_paragraph(pytanie1.Label, style='Table Header')

    #paragraph.paragraph_format.keep_with_next = True
   
    table1 = doc1.add_table(rows=3, cols=2, style='P&P_1')
    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = ''
    hdr1_cells[1].text = 'Percentage of organizations'
    hdr2_cells = table1.rows[1].cells
    hdr2_cells[0].text = 'Yes'
    hdr2_cells[1].text = pytanie1.wyniki1[0][0]
    hdr3_cells = table1.rows[2].cells
    hdr3_cells[0].text = 'No'
    hdr3_cells[1].text = pytanie1.wyniki1[0][1]
    paragraph = doc1.add_paragraph('Number of organization ' + str(pytanie1.freq[0][-1]), style='Note' )
    
    if pytanie1.rounding == True:
            paragraph = doc1.add_paragraph('the total does not equals 100%', style='Note' )
    if pytanie1.exceeds100 == True:
            paragraph = doc1.add_paragraph('the total does exeeds 100%', style='Note' )
    #paragraph=doc1.add_paragraph("/n") 
    r= doc1.add_paragraph(style ="Normal")
    #r.add_break()
    parag.paragraph_format.keep_with_next = True
    global row 
    col= 0 
    
    sheet1.write(row, col, pytanie1.Label)
    row+=1
    for a in pytanie1.wyniki1:
        sheet1.write(row, col, a)
        col+=1
        row+=1

    

def drukuj_percentyle(pytanie1):
    if pytanie1.insufficient:
        pass
    else:
        x=0
        r=0
        m=0
        n=0
        y=0
        parag = doc1.add_paragraph(pytanie1.Label, style='Table Header') #  parag.paragraph_format.keep_with_next = True

        
        
        
        table1 = doc1.add_table(rows=pytanie1.ile_grup+1, cols=6, style='P&P_1')
        hdr1_cells = table1.rows[0].cells
        hdr1_cells[0].text = ''
        hdr1_cells[1].text = '25th Percentile'  #font.supersript
        hdr1_cells[2].text = 'Median'           #font.supersript
        hdr1_cells[3].text = 'Average'          #font.supersript
        hdr1_cells[4].text = '75th Percentile'  #font.supersript
        hdr1_cells[5].text = 'No. of Responses' #font.supersript
        
        for m in range(1, pytanie1.ile_grup +1):
           # for n in range (1, 4 + 1):
                table1.rows[m].cells[5].text = str(pytanie1.wyniki1[m-1][-1])
            #    x+=1 
        for m in range(pytanie1.ile_grup):
            table1.rows[r+1].cells[0].text = pytanie1.labele_grup[r]       
            r+=1
        for m in range(pytanie1.ile_grup):
                for y in range(4):
                    table1.rows[m+1].cells[y+1].text = str(pytanie1.wyniki1[m][y])
                    #y+=1

        
        
        
        p = doc1.add_paragraph(style ="Normal")
       
        run = p.add_run()
    
def drukuj_OB(pytanie1):
    x=0
    r=0
    t=0
    m=0
    n=0
    y=0
    v=0
    parag = doc1.add_paragraph(pytanie1.Label, style='Table Header')

    
    
    
    table1 = doc1.add_table(rows=pytanie1.ile_grup+1, cols=pytanie1.ile_kolumn+2, style='P&P_1')
    for m in range(1, pytanie1.ile_grup+1):
        for n in range(1, pytanie1.ile_kolumn+1):
            table1.rows[m].cells[n].text = pytanie1.wyniki1[m-1][n-1]
            x+=1 
            v+=1
    for m in range(pytanie1.ile_grup):
        table1.rows[r+1].cells[0].text = pytanie1.labele_grup[r]       
        r+=1
    for m in range(pytanie1.ile_kolumn):    
        table1.rows[0].cells[t+1].text = pytanie1.referenceTableLabelssplit[t]
        t+=1
        table1.rows[0].cells[t+1].text = "No of responses:"
    for m in range(pytanie1.ile_grup):
        table1.rows[y+1].cells[pytanie1.ile_kolumn+1].text = str(pytanie1.freq[y][-1])
        y+=1
    for x in table1.rows:
        x.AllowBreakAcrossPages = False
    


    if pytanie1.rounding == True:
        paragraph = doc1.add_paragraph('the total does not equals 100%', style='Note' )

#    if pytanie1.exceeds100 == True:
#        paragraph = doc1.add_paragraph('the total does exeeds 100%', style='Note' )
        
    if pytanie1.example == True:    
        paragraph = doc1.add_paragraph('notka przykładowa', style='Note' )
    
    paragraph=doc1.add_paragraph(style ="Normal")
    parag.paragraph_format.keep_with_next = True
    r = paragraph.add_run()
    #r.add_break()
    i=0
    global row 
#    for m in range(pytanie1.ile_kolumn):    
#        sheet1.write(row, m+1, pytanie1.referenceTableLabelssplit[t])
#       
#        sheet1.write(row, m+1, "No of responses:")
#    row +=1
#    for m in range(pytanie1.ile_grup):
#        sheet1.write(r, 0,  pytanie1.labele_grup[r])      
#        row+=1
#        r+=1
#    for m in range(1, pytanie1.ile_grup+1):
#        for n in range(1, pytanie1.ile_kolumn+1):
#            table1.rows[m].cells[n].text = pytanie1.wyniki1[m-1][n-1]
#            x+=1 
#            v+=1


#    col =0
#    sheet1.write(row, col, pytanie1.Label)
#    row+=1
#    for a in pytanie1.referenceTableLabelssplit:
#        sheet1.write(row, col, a)  
#        col+=1
#    row+=1    
#    for a in pytanie1.labele_grup:
#        col1=0
#        sheet1.write(row, col1, a)
#        col1+=1
#        col2=0
#        for a in pytanie1.wyniki1:
#            sheet1.write(row, col2, a)
#            col2+=1
#            i+=1    
#            row+=1

def drukuj_CB_EG(pytanie1):
    x=0
    r=0
    t=0
    m=0
    n=0
    y=0
    paragraph = doc1.add_paragraph(pytanie1.label, style='Table Header')
    table1 = doc1.add_table(rows=pytanie1.ile_grup+1, cols=pytanie1.ile_kolumn+2, style='P&P_1')
    for m in range(1, pytanie1.ile_grup+1):
        for n in range(1, pytanie1.ile_kolumn+1):
            table1.rows[m].cells[n].text = pytanie1.tabelaprocenty[x]
            x+=1        
    for m in range(pytanie1.ile_grup):
        table1.rows[r+1].cells[0].text = pytanie1.grupy[r]       
        r+=1
    for m in range(pytanie1.ile_kolumn):    
        table1.rows[0].cells[t+1].text = pytanie1.kolumny[t]
        t+=1
        table1.rows[0].cells[t+1].text = "No of responses:"
    for m in range(pytanie1.ile_grup):
        table1.rows[y+1].cells[pytanie1.ile_kolumn].text = str(pytanie1.tabelaprocenty[y]) #(pytanie1.odpcount[y])
        y+=1
    if pytanie1.rounding == True:
        paragraph = doc1.add_paragraph('the total does not equals 100%', style='Note' )
    
    if pytanie1.example == True:    
        paragraph = doc1.add_paragraph('notka przykładowa', style='Note' )
    
#    r = paragraph.add_run()
#    r.add_break()
    p = doc1.add_paragraph(style ="Normal")
    run = p.add_run()
    
    
def drukuj_CB_NEG(pytanie1):
    x=0
    r=0
    t=0
    m=0
    n=0
    y=0
    paragraph = doc1.add_paragraph(pytanie1.label, style='Table Header').keep_with_next = True
    table1 = doc1.add_table(rows=pytanie1.ile_grup+1, cols=2, style='P&P_1')
    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = '' # custom label
    hdr1_cells[1].text = 'Percentage of organizations'
    for m in range(pytanie1.ile_grup):
        table1.rows[r+1].cells[0].text = pytanie1.grupy[r]       
        table1.rows[r+1].cells[1].text = str(pytanie1.tabelaprocenty[r])
        r+=1
#    for m in range(1, pytanie1.ile_grup):
#        table1.rows[y+1].cells[pytanie1.ile_kolumn].text = str(pytanie1.tabelaprocenty[y]) #(pytanie1.odpcount[y])
#        y+=1 
    
    
    paragraph = doc1.add_paragraph('Number of organization ' + str(pytanie1.number), style='Note' )
    #paragraph=doc1.add_paragraph("/n") 
#    r= paragraph.add_run()
#    r.add_break()
    p = doc1.add_paragraph(style ="Normal")
    run = p.add_run()


def drukuj_Procenty_OB_EG_Elig(pytanie1):
    x=0
    r=0
    t=0
    m=0
    n=0
    y=0
    paragraph = doc1.add_paragraph(pytanie1.Label, style='Table Header')
    table1 = doc1.add_table(rows=pytanie1.ile_grup+1, cols=3, style='P&P_1')
    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = '' # custom label
    hdr1_cells[1].text = 'Percentage of organizations'
    hdr1_cells[2].text = 'No. of responses'
    for m in range(pytanie1.ile_grup):
        table1.rows[r+1].cells[0].text = pytanie1.grupy[r]       
        table1.rows[r+1].cells[1].text = str(pytanie1.tabelaprocenty[r])
        table1.rows[r+1].cells[2].text = str(pytanie1.odp[r])
        r+=1
#    for m in range(pytanie1.ile_grup):
#        table1.rows[y+1].cells[pytanie1.ile_kolumn].text = str(pytanie1.tabelaprocenty[y]) #(pytanie1.odpcount[y])
#        y+=1 
#    for m in range(pytanie1.ile_grup):
#        table1.rows[t+1].cells[pytanie1.ile_kolumn+1].text = str(pytanie1.odp[t])
#        t+=1
    
    #paragraph = doc1.add_paragraph('Number of organization ' + str(pytanie1.number), style='Note' )
   # body.append(paragraph(""))
      
    doc1.add_paragraph(style ="Normal")
    #paragraph.add_run("")
        #run.add_break()
#for x in ex.tab:
#        if type(x) is ex.Cb and ex.Cb.zakres[0][3:6] == "NEG" and ex.Cb.zakres[0][-3:] == "Yes":
#            ex.drukuj_main(x)
#        if type(x) is ex.Cb and ex.Cb.zakres[0][3:5] == "EG" and ex.Cb.zakres[0][-3:] == "Yes":
#            ex.drukuj_Procenty_OB_EG_Elig(x)     
#        if type(x) is ex.Cb and ex.Cb.zakres[0][3:6] == "NEG":
#            ex.drukuj_CB_NEG(x)
#        if type(x) is ex.Cb:
#            ex.drukuj_OB(x)              
#        if type(x) is ex.Percentiles:
#            ex.drukuj_percentyle(x) 
#        if type(x) is ex.Cb and :
#            ex.drukuj_CB_EG(x)
#   
ins=[]
brak=[]
for x in ex.tab:
        if x.ile_grup == 0 and x.mozna == True and x.insufficient == False:
            drukuj_main(x)
        elif x.Type_of_question == "radio_buttons" and x.mozna == True and x.insufficient == False:
            drukuj_OB(x)     
        elif x.Type_of_question == "checkboxes" and x.mozna == True and x.insufficient == False:
            drukuj_OB(x)              
        elif (x.Type_of_question == "double" or x.Type_of_question == "percentage") and x.mozna == True and x.insufficient == False:
            drukuj_percentyle(x)        
        elif x.mozna == False: 
           brak.append(x.QBcode)     
        elif x.insufficient == True:
           ins.append(x.QBcode)  

file1 = open("ins.txt","w")#write mode 
for x in brak:
    file1.write(x)       
    file1.write("\n") 
file1.close() 

file2 = open("insufficient.txt","w")#write mode 
for x in ins:
    file2.write(x)       
    file2.write("\n")  
file2.close() 




#book.add_sheet('Sheet 2')
#sheet1.write(0,0,'A1')
#sheet1.write(0,1,'B1')
#row1 = sheet1.row(1)
#row1.write(0,'A2')
#row1.write(1,'B2')
#sheet1.col(0).width = 10000
#sheet2 = book.get_sheet(1)
#sheet2.row(0).write(0,'Sheet 2 A1')
#sheet2.row(0).write(1,'Sheet 2 B1')
#sheet2.flush_row_data()
#sheet2.write(1,0,'Sheet 2 A3')
#sheet2.col(0).width = 5000
#sheet2.col(0).hidden = True
book.save('simple.xls')



doc1.save('pp.docx')         
print("--- %s seconds ---" % (time.time() - start_time))       